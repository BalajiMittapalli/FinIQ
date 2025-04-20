import subprocess
from docx import Document
from docx.shared import Inches
import os

def auto_draft_response(metadata):
    # Use Ollama to draft the response
    prompt = f"""You are a professional tax‑law assistant drafting formal reply letters to Income‑Tax Department notices.

    Using the information below, compose a clear, concise, and courteous response letter. The tone should be respectful but firm, addressing each point raised in the Show‑Cause Notice. Include:

      • A proper salutation and reference to the Notice DIN
      • A brief introduction of the taxpayer and their PAN/GSTIN
      • A point‑by‑point rebuttal or explanation for each allegation (e.g., source of cash deposit)
      • Citation of supporting documents or evidence where relevant
      • A polite closing asking for confirmation of receipt and any further clarifications
      • Date and place of signing

---
Given the following extracted text from an Income Tax Notice, please draft a formal reply letter. Extract all the necessary information from the text.
{metadata}
---
Draft the letter below this line:
    """
    try:
        result = subprocess.run(['ollama', 'run', 'mistral', prompt], capture_output=True, text=True, check=True, encoding='utf-8')
        response_letter = result.stdout
    except subprocess.CalledProcessError as e:
        print(f"Error running Ollama: {e}")
        response_letter = "Error generating response."
    return response_letter

def save_response_to_docx(response_letter, output_path="response.docx"):
    if os.path.exists(output_path):
        os.remove(output_path)
    document = Document()
    document.add_heading('Response Letter', 0)
    document.add_paragraph(response_letter)
    document.save(output_path)
    print(f"Response letter saved to {output_path}")